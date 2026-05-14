from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/admin/Pg-web-Prototype/Aggregation Editor")
PRD_DIR = ROOT / "prd"
ASSET_DIR = PRD_DIR / "assets" / "login-registration-flow-v2"
EXCEPTION_ASSET_DIR = PRD_DIR / "assets" / "login-registration-flow-v2-exceptions"
OUT_PATH = PRD_DIR / "login-registration-flow-v2-prd.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = doc.styles[f"Heading {level}"]
    run = p.add_run(text)
    run.font.color.rgb = RGBColor(0x11, 0x24, 0x40)
    return p


def add_body(doc, text):
    p = doc.add_paragraph(style="Body Text")
    p.add_run(text)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_image_block(doc, title, filename, note):
    add_heading(doc, title, level=3)
    add_body(doc, note)
    img_path = ASSET_DIR / filename
    if img_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(img_path), width=Inches(6.4))
        caption = doc.add_paragraph(style="Caption")
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.add_run(filename.replace(".png", "").replace("-", " ").title())


def add_exception_image_block(doc, title, filename, note):
    add_heading(doc, title, level=3)
    add_body(doc, note)
    img_path = EXCEPTION_ASSET_DIR / filename
    if img_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(img_path), width=Inches(6.4))
        caption = doc.add_paragraph(style="Caption")
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.add_run(filename.replace(".png", "").replace("-", " ").title())


def build_doc():
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)
    styles["Body Text"].font.name = "Arial"
    styles["Body Text"].font.size = Pt(10.5)
    styles["Caption"].font.name = "Arial"
    styles["Caption"].font.size = Pt(9)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("登录注册流程 PRD V2")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x11, 0x24, 0x40)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("邮箱先输入，再按账号状态分流")
    subtitle_run.font.size = Pt(11)
    subtitle_run.font.color.rgb = RGBColor(0x5A, 0x6B, 0x85)

    meta = doc.add_table(rows=2, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.style = "Table Grid"
    meta.cell(0, 0).text = "文档版本"
    meta.cell(0, 1).text = "V2"
    meta.cell(1, 0).text = "适用范围"
    meta.cell(1, 1).text = "首页登录注册弹窗"
    for r in meta.rows:
      for c in r.cells:
        for p in c.paragraphs:
          p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_shading(meta.cell(0, 0), "EAF4FF")
    set_cell_shading(meta.cell(1, 0), "EAF4FF")

    add_heading(doc, "1. 文档目标")
    add_bullets(
        doc,
        [
            "说明当前站点认证弹窗的完整流程与分流规则。",
            "明确“邮箱先输入，再判断账号状态”的新方案。",
            "为后续接入真实账号系统、邮箱验证与 OAuth 提供清晰产品依据。",
            "配套最新页面截图，便于设计、前端、后端同步。",
        ],
    )

    add_heading(doc, "2. 方案概述")
    add_body(
        doc,
        "本版流程采用“先选方式，再录邮箱，再按账号状态分流”的结构。用户进入邮箱链路后，系统先判断邮箱属于已注册、未激活还是新用户，再分别引导到登录、邮箱验证或注册补全页面。"
    )
    add_numbered(
        doc,
        [
            "点击首页导航栏“登录”。",
            "先展示登录方式选择页。",
            "选择“使用邮箱继续”后，先进入邮箱输入页。",
            "根据邮箱状态分流到密码登录、邮箱验证提示或注册补全页。",
        ],
    )

    add_heading(doc, "3. 流程总览")
    add_bullets(
        doc,
        [
            "登录方式选择",
            "邮箱输入",
            "已注册用户 -> 密码登录",
            "未激活用户 -> 邮箱验证提示",
            "新用户 -> 注册补全",
            "忘记密码",
        ],
    )

    add_heading(doc, "4. 账号状态规则")
    state_table = doc.add_table(rows=1, cols=3)
    state_table.style = "Table Grid"
    hdr = state_table.rows[0].cells
    hdr[0].text = "状态"
    hdr[1].text = "说明"
    hdr[2].text = "下一步"
    for cell in hdr:
        set_cell_shading(cell, "D9EEF7")
    rows = [
        ("existing_active", "邮箱已注册且已完成验证", "进入密码登录页"),
        ("existing_pending_activation", "邮箱已注册但未完成验证", "进入邮箱验证提示页"),
        ("new_user", "邮箱不存在，需要新建账号", "进入注册补全页"),
    ]
    for a, b, c in rows:
        cells = state_table.add_row().cells
        cells[0].text = a
        cells[1].text = b
        cells[2].text = c

    add_heading(doc, "5. 页面说明与截图")
    add_image_block(
        doc,
        "5.1 登录方式选择页",
        "01-method-selection.png",
        "目标：降低首屏复杂度，让用户先选认证方式。页面提供 Google、Apple、邮箱三种入口。",
    )
    add_image_block(
        doc,
        "5.2 邮箱输入页",
        "02-email-entry.png",
        "目标：先获取邮箱，避免用户在未知账号状态下就输入密码。提交后根据邮箱状态分流。",
    )
    add_image_block(
        doc,
        "5.3 已注册用户密码登录页",
        "03-existing-user-login.png",
        "目标：仅对已注册邮箱展示密码输入，减少无效输入。页面支持忘记密码和更换邮箱。",
    )
    add_image_block(
        doc,
        "5.4 未激活邮箱提示页",
        "04-pending-activation.png",
        "目标：明确提示用户先完成邮箱验证。页面提供重新发送邮件、更换邮箱、返回登录三类动作。",
    )
    add_image_block(
        doc,
        "5.5 新用户注册补全页",
        "05-new-user-register.png",
        "目标：新邮箱直接进入开户注册，只要求用户设置密码和确认密码。",
    )
    add_image_block(
        doc,
        "5.6 忘记密码页",
        "06-reset-password.png",
        "目标：为已注册用户提供找回密码入口，提交后由邮件服务发送重设链接。",
    )

    add_heading(doc, "6. 异常态与提示页")
    add_exception_image_block(
        doc,
        "6.1 第三方登录占位提示",
        "01-google-oauth-placeholder.png",
        "目标：在尚未接入真实 OAuth 的阶段，明确告诉评审与开发这是待接入能力，而不是页面无响应。",
    )
    add_exception_image_block(
        doc,
        "6.2 邮箱输入校验失败",
        "02-email-entry-invalid.png",
        "目标：当用户未输入邮箱或邮箱格式非法时，阻断后续分流，并给出即时修正提示。",
    )
    add_exception_image_block(
        doc,
        "6.3 未激活邮箱返回登录后的提示",
        "03-unactivated-login-warning.png",
        "目标：用户从“查收邮箱完成验证”页返回登录后，仍能明确知道当前邮箱尚未完成验证。",
    )
    add_exception_image_block(
        doc,
        "6.4 注册确认密码不一致",
        "04-register-password-mismatch.png",
        "目标：用户在注册补全页两次输入密码不一致时，阻断注册并提示修正确认密码。",
    )
    add_exception_image_block(
        doc,
        "6.5 重新发送激活邮件成功提示",
        "05-resend-activation-feedback.png",
        "目标：用户在未激活状态下触发“重新发送邮件”后，获得明确反馈。",
    )
    add_exception_image_block(
        doc,
        "6.6 重设密码提交成功提示",
        "06-reset-password-success.png",
        "目标：用户在忘记密码页提交邮箱后，明确获知邮件已发出或准备发出。",
    )

    add_heading(doc, "7. 关键交互规则")
    add_bullets(
        doc,
        [
            "不要求用户在未知账号状态下先输入密码。",
            "新老用户与未激活用户的分流尽量自动完成。",
            "注册后的邮箱验证是必要步骤，不应默认绕过。",
            "所有阻断状态都必须给出明确下一步操作。",
            "弹窗尺寸固定，不随状态变化而跳动。",
        ],
    )

    add_heading(doc, "8. 当前实现边界")
    add_bullets(
        doc,
        [
            "Google / Apple 登录入口为前端预留入口，尚未接入真实 OAuth。",
            "邮箱状态判断当前使用 mock 数据集合。",
            "登录成功、注册成功、邮箱验证成功、重设密码成功均为前端原型反馈。",
            "生产环境中，账号状态与认证结果必须由服务端接口决定。",
        ],
    )

    add_heading(doc, "9. 后续接入建议")
    add_numbered(
        doc,
        [
            "接入邮箱状态查询接口，返回已注册 / 未激活 / 新用户三态。",
            "接入 Google OAuth。",
            "接入 Apple OAuth。",
            "接入邮箱密码登录接口。",
            "接入邮箱注册接口。",
            "接入邮箱激活邮件发送接口。",
            "接入忘记密码邮件发送接口。",
            "统一前后端密码规则与错误提示文案。",
        ],
    )

    add_heading(doc, "10. 验收标准")
    add_numbered(
        doc,
        [
            "点击登录后，先展示登录方式选择页。",
            "点击使用邮箱继续后，先进入邮箱输入页。",
            "已注册邮箱提交后进入密码登录页。",
            "未激活邮箱提交后进入邮箱验证提示页。",
            "新邮箱提交后进入注册补全页。",
            "注册成功后进入邮箱验证提示页。",
            "忘记密码入口可从密码登录页进入。",
            "各页面支持正确返回，不出现死循环或错误跳转。",
            "弹窗尺寸在不同状态下保持一致。",
            "每个页面状态都有对应截图可供对照。",
        ],
    )

    doc.save(OUT_PATH)


if __name__ == "__main__":
    build_doc()
